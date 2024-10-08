from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from Utils import Format_Ctrl_Utils
from Language_Selection import C_Compiler, CPP_Compiler
from docx.shared import Inches
from .CmdSS import ScreenShotOutput

class WordDocument_Handle:
    def __init__(self, doc: Document, befIpOpList: list, genericFont: list):
        # Generic Font indicates [fontname: str, fontsize: int, bold: bool]
        self.document = doc
        self.Utils = Format_Ctrl_Utils()
        self.befIpOp = befIpOpList
        self.GenericFont = genericFont
      


    def set_headFoot(self, hfList: list):
        """
        ## Adding Header and Footer
        """

        section = self.document.sections[0]
        header = section.header
        HeadPara = header.paragraphs[0]
        HeadPara.text = f"\t{hfList[0]}" # Insert header


        section2 = self.document.sections[0]
        header2 = section.footer
        HeadPara2 = header2.paragraphs[0]
        HeadPara2.text = f"\t{hfList[1]}" # Insert Footer


    def set_question(self, question: str):
        """
        ## Sets Question First
        """
        para = self.document.add_paragraph("")
        setter = para.add_run(question).font
        setter.name = self.GenericFont[0]
        setter.size = Pt(int(self.GenericFont[1]))
        setter.bold = True

    
    def set_details(self, DetailsList: list, fontstyle: str, fontsize: int, isbold: bool):
        """
        ## Set Details and fonts
        """
        para = self.document.add_paragraph("")
        detail_str = self.Utils.person_details_formatter(detailsList=DetailsList)

        # Font setter
        setter = para.add_run(detail_str).font
        setter.name = fontstyle
        setter.size = Pt(int(fontsize))
        setter.bold = isbold

    
    def set_befCode(self, fontsize, isbold: bool):
        """
        # Sets Before Code Title
        
        """
        para = self.document.add_paragraph("")
        setter = para.add_run(f"//{self.befIpOp[0]}:").font
        setter.name = self.GenericFont[0]
        setter.size = Pt(int(fontsize))
        setter.bold = isbold

    def set_befOutput(self, fontsize, isbold: bool):
        """
        
        # Sets Before Ouput Title
        """
        para = self.document.add_paragraph("")
        setter = para.add_run(f"{self.befIpOp[1]}:").font
        setter.name = self.GenericFont[0]
        setter.size = Pt(int(fontsize))
        setter.bold = isbold
        


    def set_code(self, address: str):
        code = ""
        # Read code from C file
        with open(address, "r") as ip:
            code = ip.read()
        
        # Sets Before Ouput Title
        para = self.document.add_paragraph("")
        setter = para.add_run(code).font
        setter.name = self.GenericFont[0]
        setter.size = Pt(int(self.GenericFont[1]))
        setter.bold = self.GenericFont[2]
        self.document.add_page_break()


    def set_outputCases_text(self, Address, frequency, inputStr, fontsize, isbold: bool, mode) ->str:
        """
        # Sets output (Text output)
        
        """
        count = int(frequency)
        if mode == "c":
            op = C_Compiler(address=Address, frequency_input=count, input_str=inputStr)

        elif mode == "cpp":
            op = CPP_Compiler(address=Address, frequency_input=count, input_str=inputStr)

        output = op.get_output()
        para = self.document.add_paragraph("")
        setter = para.add_run(output).font
        setter.name = self.GenericFont[0]
        setter.size = Pt(int(fontsize))
        setter.bold = isbold
        self.document.add_page_break()

    def get_details(self):
        """
        ## Get user details, names etc
        returns list [name, roll, section, course]
        """
        path = "UI\\Features\\Data\\details.txt"
        lst = []
        with open(path, "r") as fp:
            lst = fp.read().split("#")
        
        return lst

    def set_outputCases_screenShot(self, Address, frequency, inputStr, fontsize, isbold: bool, mode):
        """
        Sets output (Screen shot - Picture)
        """
        count = int(frequency)
        if mode == "c":
            op = C_Compiler(address=Address, frequency_input=count, input_str=inputStr)

        elif mode == "cpp":
            op = CPP_Compiler(address=Address, frequency_input=count, input_str=inputStr)

        output = op.get_output()

        detail_lst = self.get_details()
        user_name = detail_lst[0]

        ## TODO: Get file name here
        ScreenShotOutput().create_command_image()

        self.document.add_picture('Format_Control\\command_image_final.png', width=Inches(1.25))
        self.document.add_page_break()
        

    
    def save_file(self, fileaddress):
        # To save file
        self.document.save(fileaddress)
